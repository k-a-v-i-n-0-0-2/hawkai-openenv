import docx
import json

doc = docx.Document(r'FINAL_28_STATE_DOCX.docx')

# Get state names from paragraphs
states = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
print(f"States ({len(states)}): {states}")
print(f"Tables: {len(doc.tables)}")
print()

all_data = {}
for ti, table in enumerate(doc.tables):
    state_name = states[ti] if ti < len(states) else f"Unknown_{ti}"
    rows = []
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    all_data[state_name] = rows
    
    # Print header + first 5 data rows
    print(f"=== {state_name} (TABLE {ti}, {len(table.rows)} rows) ===")
    for r in rows[:6]:
        print("  ", r)
    if len(rows) > 6:
        print(f"  ... ({len(rows)} total rows)")
    print()

# Save to json for further processing
with open("docx_extracted.json", "w", encoding="utf-8") as f:
    json.dump(all_data, f, indent=2, ensure_ascii=False)
print("Saved to docx_extracted.json")
